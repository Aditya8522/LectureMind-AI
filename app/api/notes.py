"""
app/api/notes.py
─────────────────
Notes generation API endpoints.

Endpoints:
  POST /api/notes                        → generate notes for a video (summary or detailed)
  GET  /api/notes/{video_id}             → get the most recent notes for a video
  POST /api/notes/download               → download notes as PDF or DOCX
"""

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from sqlalchemy.orm import Session
import io
import re

from app.models.db import get_db, Video, Note
from app.core.retrieval import retrieve_all_chunks_for_video
from app.core.llm import generate_notes

router = APIRouter(prefix="/api/notes", tags=["notes"])


# ── Request / Response Schemas ─────────────────────────────────────────────────

class NotesRequest(BaseModel):
    video_id: str            # internal SQLite video row ID
    mode: str = "summary"    # "summary" | "detailed"


class NotesResponse(BaseModel):
    video_id: str
    mode: str
    content: str             # Markdown string


class DownloadRequest(BaseModel):
    content: str             # Markdown content
    title: str = "Lecture Notes"
    format: str = "pdf"      # "pdf" | "docx"


# ── Markdown Parser ────────────────────────────────────────────────────────────

def _parse_md_lines(md: str) -> list:
    """
    Parse Markdown into a list of (tag, data) tuples for document rendering.
    tags:
      - ('h1'|'h2'|'h3'|'h4', str)
      - ('bullet'|'numbered', str)
      - ('blockquote', str)
      - ('table', list[list[str]])
      - ('code', str)
      - ('hr', '')
      - ('text', str)
      - ('blank', '')
    """
    lines = md.split("\n")
    result = []
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        if table_buf:
            rows = []
            for tline in table_buf:
                s = tline.strip()
                if re.match(r"^[|\s\-:]+$", s):
                    continue  # Separator line
                cells = [c.strip() for c in s.strip("|").split("|")]
                if cells and any(cells):
                    rows.append(cells)
            if rows:
                result.append(("table", rows))
            table_buf.clear()

    for line in lines:
        if line.startswith("```"):
            flush_table()
            if in_code:
                result.append(("code", "\n".join(code_buf)))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        stripped = line.rstrip()

        # Table detection
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buf.append(stripped)
            continue
        else:
            flush_table()

        if stripped.startswith("#### "):
            result.append(("h4", stripped[5:]))
        elif stripped.startswith("### "):
            result.append(("h3", stripped[4:]))
        elif stripped.startswith("## "):
            result.append(("h2", stripped[3:]))
        elif stripped.startswith("# "):
            result.append(("h1", stripped[2:]))
        elif stripped.startswith("> "):
            result.append(("blockquote", stripped[2:]))
        elif re.match(r"^\d+\.\s+", stripped):
            result.append(("numbered", re.sub(r"^\d+\.\s+", "", stripped)))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            result.append(("bullet", stripped[2:]))
        elif stripped.startswith("---"):
            result.append(("hr", ""))
        elif stripped == "":
            result.append(("blank", ""))
        else:
            result.append(("text", stripped))

    flush_table()
    return result


def _strip_md(text: str) -> str:
    """Strip basic markdown formatting for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(\d+:\d+(?::\d+)?)\]", r"[\1]", text)
    return text


def _clean_for_pdf(text: str) -> str:
    """Strip emojis and normalize special Unicode quotes/dashes for FPDF core Helvetica font."""
    text = _strip_md(text)
    # Remove 4-byte UTF-8 emojis
    emoji_pattern = re.compile(r"[\U00010000-\U0010ffff]", flags=re.UNICODE)
    text = emoji_pattern.sub("", text)
    # Common Unicode substitutions
    text = (
        text.replace("\u2014", " -- ")
        .replace("\u2013", "-")
        .replace("\u2022", "*")
        .replace("\u2019", "'")
        .replace("\u2018", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2192", "->")
        .replace("\u2190", "<-")
        .replace("\u2026", "...")
    )
    # Filter to latin-1 safe characters
    return text.encode("latin-1", "ignore").decode("latin-1").strip()


# ── DOCX Generator ────────────────────────────────────────────────────────────

def _build_docx(content: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def set_para_font(para, size_pt, bold=False, color=None, italic=False):
        for run in para.runs:
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)

    def add_inline_formatted(para, text):
        """Add text with **bold** and `code` inline formatting."""
        parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = para.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
            else:
                para.add_run(part)

    parsed = _parse_md_lines(content)

    for tag, data in parsed:
        if tag == "h1":
            p = doc.add_heading(level=1)
            p.clear()
            add_inline_formatted(p, _strip_md(data))
            set_para_font(p, 18, bold=True, color=(15, 23, 42))

        elif tag == "h2":
            p = doc.add_heading(level=2)
            p.clear()
            add_inline_formatted(p, _strip_md(data))
            set_para_font(p, 14, bold=True, color=(30, 41, 59))

        elif tag == "h3":
            p = doc.add_heading(level=3)
            p.clear()
            add_inline_formatted(p, _strip_md(data))
            set_para_font(p, 12, bold=True, color=(51, 65, 85))

        elif tag == "h4":
            p = doc.add_paragraph()
            add_inline_formatted(p, _strip_md(data))
            set_para_font(p, 11, bold=True, italic=True, color=(71, 85, 105))

        elif tag == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            add_inline_formatted(p, data)
            set_para_font(p, 10.5, color=(30, 41, 59))

        elif tag == "numbered":
            p = doc.add_paragraph(style="List Number")
            p.clear()
            add_inline_formatted(p, data)
            set_para_font(p, 10.5, color=(30, 41, 59))

        elif tag == "blockquote":
            p = doc.add_paragraph()
            add_inline_formatted(p, data)
            set_para_font(p, 10.5, italic=True, color=(67, 56, 202))
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "720")
            pPr.append(ind)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "EEF2FF")
            pPr.append(shd)

        elif tag == "table":
            rows = data
            if rows:
                num_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=num_cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(num_cols):
                        cell_val = row[c_idx] if c_idx < len(row) else ""
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = _strip_md(cell_val)
                        if r_idx == 0:
                            for p in cell.paragraphs:
                                set_para_font(p, 10, bold=True, color=(15, 23, 42))
                            tcPr = cell._tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:val"), "clear")
                            shd.set(qn("w:color"), "auto")
                            shd.set(qn("w:fill"), "F1F5F9")
                            tcPr.append(shd)
                        else:
                            for p in cell.paragraphs:
                                set_para_font(p, 9.5, color=(51, 65, 85))
                doc.add_paragraph("")

        elif tag == "code":
            p = doc.add_paragraph()
            run = p.add_run(data)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(15, 23, 42)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F1F5F9")
            pPr.append(shd)

        elif tag == "hr":
            p = doc.add_paragraph("─" * 60)
            p.runs[0].font.color.rgb = RGBColor(200, 200, 200)

        elif tag == "text":
            p = doc.add_paragraph()
            add_inline_formatted(p, data)
            set_para_font(p, 10.5, color=(30, 41, 59))

        elif tag == "blank":
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── PDF Generator ─────────────────────────────────────────────────────────────

def _build_pdf(content: str, title: str) -> bytes:
    from fpdf import FPDF

    safe_hdr_title = _clean_for_pdf(title)[:75]

    class NotePDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 8)
            self.set_text_color(100, 116, 139)
            self.cell(0, 7, safe_hdr_title, align="L")
            self.ln(1)
            self.set_draw_color(226, 232, 240)
            self.line(self.get_x(), self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(148, 163, 184)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = NotePDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(left=18, top=18, right=18)
    pdf.add_page()

    parsed = _parse_md_lines(content)

    for tag, data in parsed:
        if tag == "h1":
            clean = _clean_for_pdf(data)
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(15, 23, 42)
            pdf.multi_cell(0, 8, clean, new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(79, 70, 229)
            pdf.set_line_width(0.5)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(3)

        elif tag == "h2":
            clean = _clean_for_pdf(data)
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12.5)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(0, 7, clean, new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(226, 232, 240)
            pdf.set_line_width(0.3)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(2)

        elif tag == "h3":
            clean = _clean_for_pdf(data)
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(51, 65, 85)
            pdf.multi_cell(0, 6, clean, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        elif tag == "h4":
            clean = _clean_for_pdf(data)
            pdf.set_font("Helvetica", "BI", 9.5)
            pdf.set_text_color(71, 85, 105)
            pdf.multi_cell(0, 5, clean, new_x="LMARGIN", new_y="NEXT")

        elif tag == "bullet":
            clean = _clean_for_pdf(data)
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(30, 41, 59)
            pdf.set_x(pdf.l_margin + 3)
            pdf.multi_cell(0, 5.5, f"-  {clean}", new_x="LMARGIN", new_y="NEXT")

        elif tag == "numbered":
            clean = _clean_for_pdf(data)
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(30, 41, 59)
            pdf.set_x(pdf.l_margin + 3)
            pdf.multi_cell(0, 5.5, clean, new_x="LMARGIN", new_y="NEXT")

        elif tag == "blockquote":
            clean = _clean_for_pdf(data)
            pdf.set_fill_color(238, 242, 255)
            pdf.set_font("Helvetica", "I", 9.5)
            pdf.set_text_color(67, 56, 202)
            pdf.set_x(pdf.l_margin + 4)
            pdf.multi_cell(
                pdf.w - pdf.l_margin - pdf.r_margin - 4,
                5.5,
                clean,
                fill=True,
                new_x="LMARGIN",
                new_y="NEXT",
            )
            pdf.ln(1)

        elif tag == "table":
            rows = data
            if rows:
                pdf.ln(2)
                pdf.set_font("Helvetica", "", 8.5)
                try:
                    with pdf.table(line_height=5.5) as table:
                        for r_idx, row in enumerate(rows):
                            table_row = table.row()
                            for cell_val in row:
                                clean_cell = _clean_for_pdf(cell_val)
                                table_row.cell(clean_cell)
                    pdf.ln(2)
                except Exception as te:
                    for row in rows:
                        clean_row = " | ".join(_clean_for_pdf(c) for c in row)
                        pdf.multi_cell(0, 5, clean_row, new_x="LMARGIN", new_y="NEXT")

        elif tag == "code":
            clean_code = _clean_for_pdf(data)
            pdf.set_fill_color(241, 245, 249)
            pdf.set_font("Courier", "", 8)
            pdf.set_text_color(15, 23, 42)
            pdf.multi_cell(0, 4.5, clean_code, fill=True, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

        elif tag == "hr":
            pdf.ln(2)
            pdf.set_draw_color(226, 232, 240)
            pdf.set_line_width(0.3)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(2)

        elif tag == "text":
            clean = _clean_for_pdf(data)
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(0, 5.5, clean, new_x="LMARGIN", new_y="NEXT")

        elif tag == "blank":
            pdf.ln(1.5)

    return bytes(pdf.output())



# ── Endpoints ──────────────────────────────────────────────────────────────────

@router.post("/download")
def download_notes(request: DownloadRequest):
    """
    Convert Markdown notes to PDF or DOCX and stream as a file download.
    """
    fmt = request.format.lower()
    safe_title = re.sub(r"[^\w\s-]", "", request.title)[:60].strip()
    filename_base = safe_title.replace(" ", "_") or "lecture_notes"

    if fmt == "docx":
        try:
            data = _build_docx(request.content, request.title)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'},
        )

    elif fmt == "pdf":
        try:
            data = _build_pdf(request.content, request.title)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.pdf"'},
        )

    else:
        raise HTTPException(status_code=400, detail="format must be 'pdf' or 'docx'.")


@router.post("", response_model=NotesResponse)
def create_notes(request: NotesRequest, db: Session = Depends(get_db)):
    """
    Generate study notes from the lecture transcript.

    Steps:
      1. Validate video exists and is ready
      2. Fetch ALL chunks from ChromaDB (full transcript context)
      3. Call Gemini to generate structured Markdown notes
      4. Persist notes to SQLite
      5. Return Markdown content
    """
    # Validate mode
    if request.mode not in ("summary", "detailed"):
        raise HTTPException(
            status_code=400,
            detail="mode must be 'summary' or 'detailed'."
        )

    # Validate video
    try:
        vid_id_int = int(request.video_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="video_id must be an integer.")

    video = db.query(Video).filter(Video.id == vid_id_int).first()
    if not video:
        raise HTTPException(
            status_code=404,
            detail=f"Video '{request.video_id}' not found."
        )
    if video.status != "ready":
        raise HTTPException(
            status_code=400,
            detail=f"Video is not ready yet (status: {video.status})."
        )

    # Fetch all chunks from ChromaDB
    try:
        chunks = retrieve_all_chunks_for_video(video.youtube_video_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

    if not chunks:
        raise HTTPException(status_code=404, detail="No transcript chunks found for this video.")

    # Generate notes with Gemini
    try:
        content = generate_notes(
            chunks=chunks,
            video_title=video.title or "Lecture",
            mode=request.mode,
        )
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))

    # Persist to SQLite
    note = Note(
        video_id=vid_id_int,
        mode=request.mode,
        content=content,
    )
    db.add(note)
    db.commit()
    db.refresh(note)

    print(f"[notes] Generated {request.mode} notes for video '{video.youtube_video_id}' ({len(content)} chars)")

    return NotesResponse(
        video_id=request.video_id,
        mode=request.mode,
        content=content,
    )


@router.get("/{video_id}", response_model=Optional[NotesResponse])
def get_notes(video_id: str, mode: str = "summary", db: Session = Depends(get_db)):
    """
    Retrieve the most recently generated notes for a video.
    Pass ?mode=summary or ?mode=detailed to filter.
    """
    try:
        vid_id_int = int(video_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="video_id must be an integer.")

    note = (
        db.query(Note)
        .filter(Note.video_id == vid_id_int, Note.mode == mode)
        .order_by(Note.created_at.desc())
        .first()
    )

    if not note:
        return None

    return NotesResponse(
        video_id=video_id,
        mode=note.mode,
        content=note.content,
    )

